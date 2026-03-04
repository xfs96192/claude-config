#!/usr/bin/env python3
import argparse
import json
import os
from datetime import datetime
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


MARKERS = ("产品配置策略：", "产品运作回顾：", "后市展望：")


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    return new_para


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._p
    parent = element.getparent()
    if parent is None:
        return
    parent.remove(element)


def _find_paragraph_index_containing(doc: Document, needle: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if needle in (p.text or ""):
            return i
    raise ValueError(f"未在docx中找到段落标题（包含）：{needle}")


def _find_first_paragraph_index_starting_with(
    doc: Document, start_idx: int, end_idx: int, prefix: str
) -> int:
    for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
        if (doc.paragraphs[i].text or "").strip().startswith(prefix):
            return i
    raise ValueError(f"未在docx中找到字段段落：{prefix}（范围 {start_idx}-{end_idx}）")


def _safe_output_path(template_path: str) -> str:
    base, ext = os.path.splitext(template_path)
    return f"{base}_已更新{ext}"


def _backup_inplace(path: str) -> str:
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup_path = f"{path}.bak.{ts}"
    with open(path, "rb") as src, open(backup_path, "wb") as dst:
        dst.write(src.read())
    return backup_path


def _get_continuations(product_data: Dict) -> Optional[List[str]]:
    cont = product_data.get("后市展望_续段")
    if cont is None:
        return None
    if not isinstance(cont, list) or not all(isinstance(x, str) for x in cont):
        raise ValueError("后市展望_续段 必须是字符串数组")
    return cont


def _process_product(
    doc: Document,
    product_key: str,
    product_data: Dict,
    header_idx: int,
    end_idx: int,
) -> None:
    for marker in MARKERS:
        if marker == "后市展望：" and ("后市展望" in product_data):
            pass
        if marker == "产品配置策略：" and ("产品配置策略" not in product_data):
            raise ValueError(f"{product_key}: 缺少字段 产品配置策略")
        if marker == "产品运作回顾：" and ("产品运作回顾" not in product_data):
            raise ValueError(f"{product_key}: 缺少字段 产品运作回顾")
        if marker == "后市展望：" and ("后市展望" not in product_data):
            raise ValueError(f"{product_key}: 缺少字段 后市展望")

        target_idx = _find_first_paragraph_index_starting_with(doc, header_idx, end_idx, marker)

        if marker == "产品配置策略：":
            doc.paragraphs[target_idx].text = f"{marker}{product_data['产品配置策略']}"
        elif marker == "产品运作回顾：":
            doc.paragraphs[target_idx].text = f"{marker}{product_data['产品运作回顾']}"
        elif marker == "后市展望：":
            doc.paragraphs[target_idx].text = f"{marker}{product_data['后市展望']}"

            cont = _get_continuations(product_data)
            if cont is None:
                continue

            # Replace everything after “后市展望：” within the product block, then insert new continuations.
            to_delete: List[Paragraph] = []
            for i in range(target_idx + 1, min(end_idx, len(doc.paragraphs))):
                to_delete.append(doc.paragraphs[i])

            for p in to_delete:
                _delete_paragraph(p)

            cursor = doc.paragraphs[target_idx]
            for t in cont:
                cursor = _insert_paragraph_after(cursor, t)

            _insert_paragraph_after(cursor, "")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="按产品标题子串定位，回填三只产品的【产品配置策略/产品运作回顾/后市展望】到docx模板。"
    )
    parser.add_argument("--template", required=True, help="输入docx模板路径")
    parser.add_argument("--content", required=True, help="内容JSON路径（见SKILL.md示例结构）")
    parser.add_argument(
        "--output",
        default=None,
        help="输出docx路径（默认：在模板同目录生成“*_已更新.docx”）",
    )
    parser.add_argument(
        "--inplace",
        action="store_true",
        help="直接覆盖模板文件（会自动生成 .bak.YYYYMMDD-HHMMSS 备份）",
    )
    args = parser.parse_args()

    with open(args.content, "r", encoding="utf-8") as f:
        content: Dict[str, Dict] = json.load(f)

    if not isinstance(content, dict) or not content:
        raise ValueError("内容JSON必须是非空对象，键为产品名子串")

    template_path = args.template
    if args.inplace:
        output_path = template_path
    else:
        output_path = args.output or _safe_output_path(template_path)

    if args.inplace:
        backup_path = _backup_inplace(template_path)
        print(f"[OK] Backup created: {backup_path}")

    doc = Document(template_path)

    header_positions: List[Tuple[int, str]] = []
    for product_key in content.keys():
        idx = _find_paragraph_index_containing(doc, product_key)
        header_positions.append((idx, product_key))

    header_positions.sort(key=lambda x: x[0])
    positions_by_key = {k: i for i, (idx, k) in enumerate(header_positions)}

    # Process from bottom to top to keep earlier indices stable.
    for header_idx, product_key in reversed(header_positions):
        pos = positions_by_key[product_key]
        if pos == len(header_positions) - 1:
            end_idx = len(doc.paragraphs)
        else:
            end_idx = header_positions[pos + 1][0]

        _process_product(doc, product_key, content[product_key], header_idx, end_idx)

    doc.save(output_path)
    print(f"[OK] Wrote: {output_path}")


if __name__ == "__main__":
    main()

